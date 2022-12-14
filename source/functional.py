# -*- coding: utf-8 -*-
# Файл с пользовательскими функциями для очищения от таковых файла main
from PyQt6.QtSql import *
import os
from docxtpl import DocxTemplate
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import comtypes.client
from docxtpl import DocxTemplate



db_name = "DataBaseExhibitions.db"  # имя базы данных

def db_connect():
    """Функция для подключения к базе данных.
    Параметров не принимает, база данных является константой"""
    db = QSqlDatabase.addDatabase('QSQLITE')
    db.setDatabaseName(db_name)
    if not db.open():
        return False
    else:
        print("Подключено к базе данных", db_name)
    return db

db = db_connect()    # db - глобальная переменная ссылающаяся на базу данных, создается с модуле functional.py строка 23 после функции db_connect()

def gridLayoutStartResize():
    """Начальное изменение размера сетки для дальнейшего нормального ресайза окна.
    Связано с тем, что gridLayout не получается сделать больше начального значения"""
    parameter_search = '<widget class="QWidget" name="gridLayoutWidget">'
    GUI_file = open('MainForm.ui', 'r', encoding='utf-8')
    file_info = ''
    for line in GUI_file:
        file_info += line
    inx = file_info.find(parameter_search)     # начальный индекс параметра
    inx_end = file_info.find('</rect>', inx)   # конечный индекс настройки grid layout

    setting_old = setting = file_info[inx:inx_end]
    # замена размеров для последующего уменьшения gridLayoutWidget
    file_info = file_info.replace(setting_old, replace_setting(setting))
    
    """ parameter_search = '<widget class="QWidget" name="gridLayoutWidget_2">'
    inx = file_info.find(parameter_search)     # начальный индекс параметра
    inx_end = file_info.find('</rect>', inx)   # конечный индекс настройки grid layout
    setting_old = setting = file_info[inx:inx_end] """
    # замена размеров для последующего уменьшения gridLayoutTabFiltr
    #file_info = file_info.replace(setting_old, replace_setting(setting))
    GUI_file.close()
    GUI_file = open('MainFormResize.ui', 'w', encoding='utf-8')
    GUI_file.write(file_info)
    GUI_file.close()

def replace_setting(setting):
    """Функция для сокращения части кода gridLayoutStartResize. Меняет размерности в settings"""
    inx_width_1, inx_width_2 = setting.find("<width>"), setting.find("</width>")
    setting_width = setting[inx_width_1:inx_width_2]
    setting = setting.replace(setting_width, "<width>8000")
    inx_height_1, inx_height_2 = setting.find("<height>"), setting.find("</height>")
    setting_height = setting[inx_height_1:inx_height_2]
    setting = setting.replace(setting_height, "<height>6000")
    return setting

def SQLdata_acquisition_EditWindow():
    """Получение данных с таблиц для оформления EditWindow
    Выходные данные: словарь с ключами - кодами ВУЗов и значениями - аббревиатурами ВУЗов,
    список кодов ГРНТИ, словарь с регистрационными номерами по ключам кода ВУЗа"""
    query = QSqlQuery()
    query.exec(
        #"""SELECT Код_ВУЗа, Аббревиатура, Рег_номер FROM Vyst_mo"""
        """SELECT Код_ВУЗа, Аббревиатура FROM VUZ"""
    )
    university_dict = {}
    regNum_dict = {};   # regNum_dict - словарь регистрационных номеров по ключу кода ВУЗа
    while query.next():
        current_code = query.value("Код_ВУЗа")
        university_dict.update({current_code: query.value("Аббревиатура")})  # update заменяет старое значение по ключу, если оно существует или добавляет новое
        """ if not (current_code in regNum_dict.keys()):
            regNum_dict.update({current_code: list()})
        regNum_dict[current_code].append(query.value("Рег_номер")) """
    # сортировка по ключам
    university_dict = dict(sorted(university_dict.items()))
    regNum_dict = dict(sorted(regNum_dict.items()))

    query.exec(
        """SELECT Код_рубрики FROM grntirub"""
    )
    grnti_code = []
    while query.next():
        grnti_code.append(query.value("Код_рубрики"))

    return university_dict, grnti_code, regNum_dict


def editingSQL_NIR(Edit=False, parameters_dict=False, orig_univer_code=False, orig_regNum=False):
    """Редактирование полей таблицы НИР"""
    query = QSqlQuery()
    
    if Edit:
        values = ""
        for key in parameters_dict.keys():
            if key == "Код_ВУЗа":   values += key + '=' + parameters_dict[key] + ', '
            else:   values += key + '="' + str(parameters_dict[key]) + '", '
        values = values[0:len(values)-2]
        query.exec(
            f"""UPDATE Vyst_mo SET {values} WHERE Код_ВУЗа={orig_univer_code} AND Рег_номер="{orig_regNum}" """
        )
    else:
        keys = '('
        values = "("
        for key in parameters_dict.keys():
            keys += key + ", "
            if key == "Код_ВУЗа":   values += str(parameters_dict[key]) + ', '
            else:   values += '"' + str(parameters_dict[key]) + '", '
        keys = keys[0:len(keys)-2]; keys += ')'
        values = values[0:len(values)-2] + ')'
        query.exec(
            f"""INSERT INTO Vyst_mo {keys} VALUES {values}"""
        )


def get_info_for_filtration():
    """Получение информации из таблицы для вставки её в фильтрацию"""
    query = QSqlQuery("""SELECT Код_ВУЗа, Аббревиатура FROM Vyst_mo""")
    university_dict = {}
    while query.next():
        current_code = query.value("Код_ВУЗа")
        university_dict.update({current_code: query.value("Аббревиатура")})  # update заменяет старое значение по ключу, если оно существует или добавляет новое
    # сортировка по ключам
    university_dict = dict(sorted(university_dict.items()))

    for code_uni in list(university_dict.keys()):
        query.exec(
                f"""SELECT Федеральный_округ, Город, Область FROM VUZ WHERE Код_ВУЗа={int(code_uni)}"""
            )
        query.next()
        university_dict.update({code_uni: [university_dict[code_uni], query.value("Федеральный_округ"), query.value("Город"), query.value("Область")]})
    
    query.exec(
                f"""SELECT ГРНТИ FROM Vyst_mo"""
            )

    return university_dict  #   #university_dict = {коду ВУЗа, [Аббревиатура, Федеральный округ, Город, Область]}

def get_GRNTI():
    """Возвращает все существующие в табилце Vyst_mo ГРНТИ коды в формате {"[Код_ВУЗа, Рег_номер]": "ГРНТИ"}"""
    query = QSqlQuery("""SELECT Код_ВУЗа, Рег_номер, ГРНТИ FROM Vyst_mo""")
    GRNTI_dict = {}
    while query.next():
        GRNTI_dict.update({f'[{query.value("Код_ВУЗа")}, "{query.value("Рег_номер")}"]': query.value("ГРНТИ")})
    return GRNTI_dict


def get_GRNTI_fromGRNTItable():
    """Возвращает рубрики из таблиццы grntirub"""
    query = QSqlQuery("""SELECT * FROM grntirub""")
    GRNTI_list = [] # список списков ["Код_рубрики Рубрика"]
    while query.next():
        GRNTI_list.append(query.value("Код_рубрики") + " " + query.value("Рубрика"))
    return GRNTI_list



def get_custom_table():
    """Считывание групп НИР созданных ранее"""
    if os.path.exists("data"):
        all_files = []
        for root, dirs, files in os.walk("data"):  
            for filename in files:
                all_files.append(filename[0:len(filename)-4])
        return all_files
    return []


def doc_save(path, headlines, list_values, extension="Документ Microsoft Word (*.docx)"):
    """Сохраняет данные по карточке НИР в docx или pdf файл
    path - путь
    headlines - список заголовков
    list_values - список данных для заголовков
    extension = Документ Microsoft Word (*.docx) или PDF (*.pdf)
    Размерности headlines и list_values должны совпадать"""
    if len(list_values) != len(headlines): return -1

    doc = docx.Document()

    paragraph1 = doc.add_paragraph()
    paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph1.add_run()
    run.font.size = Pt(18)
    run = paragraph1.add_run(list_values[1] + "\n")
    run.font.bold = True
    run.font.size = Pt(18)

    paragraph2 = doc.add_paragraph()
    for i in range(len(list_values)):
        if i == 2: continue
        run = paragraph2.add_run(headlines[i] + ": ")
        run.font.bold = True
        run.font.size = Pt(16)
        run = paragraph2.add_run(str(list_values[i]) + "\n")
        run.font.size = Pt(16)

    # try в случае, если файл открыт, иначе прога вылетет, по-другому не успею обработать ошибку
    try:
        if extension == "PDF (*.pdf)":
            # Сохранение в pdf
            doc.save("temp.docx") 

            doc = DocxTemplate("temp.docx")

            wdFormatPDF = 17

            in_file = os.path.abspath("temp.docx")
            out_file = os.path.abspath(path)

            word = comtypes.client.CreateObject('Word.Application')
            doc = word.Documents.Open(in_file)
            doc.SaveAs(out_file, FileFormat=wdFormatPDF)
            doc.Close()
            word.Quit()

            os.remove("temp.docx")
        else:
            # Сохранение в word
            doc.save(path[0:len(path)-5] + ".docx") 
    except: return -1
    return 0

def saveTable_toDoc(path, table_name, headlines, data, extension="Документ Microsoft Word (*.docx)"):
    """Сохраняет группу НИР в docx или pdf файл
    path - путь
    table_name - имя группы НИР
    headlines - список заголовков
    list_values - список данных для заголовков
    extension = Документ Microsoft Word (*.docx) или PDF (*.pdf)
    Размерности headlines и list_values должны совпадать"""
    if len(data[0]) != len(headlines): return -1

    # headlines = ["Аббревиатура", "Название_НИР", "Рег_номер", "ГРНТИ", "Тип", "Наличие_экспоната", "Выставки", 
    #            "Экспонат", "Научный_руководитель", "Статус_руководителя", "Код_ВУЗа"]

    doc = docx.Document()

    paragraph1 = doc.add_paragraph()
    paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph1.add_run()
    run.font.size = Pt(18)
    run = paragraph1.add_run("Группа НИР: " + table_name)
    run.font.bold = True
    run.font.size = Pt(18)

    table = doc.add_table(rows=len(data)+1, cols=len(data[0])-4)
    table.style = 'Table Grid'

    k=0
    for j in range(0, len(headlines)):
        if headlines[j] == "Код_ВУЗа" or headlines[j] == "Название_НИР" or headlines[j] == "Выставки" or headlines[j] == "Экспонат":
            continue
        cell = table.cell(0, k)
        k+=1
        cell.text = headlines[j]

    for i in range(0, len(data)):
        k=0
        for j in range(0, len(data[0])):
            if not (headlines[j] == "Код_ВУЗа" or headlines[j] == "Название_НИР" or headlines[j] == "Выставки" or headlines[j] == "Экспонат"):
                cell = table.cell(i+1, k)
                k+=1
                cell.text = data[i][j]

    # try в случае, если файл открыт, иначе прога вылетет, по-другому не успею обработать ошибку
    try:
        if extension == "PDF (*.pdf)":
            # Сохранение в pdf
            doc.save("temp.docx") 

            doc = DocxTemplate("temp.docx")

            wdFormatPDF = 17

            in_file = os.path.abspath("temp.docx")
            out_file = os.path.abspath(path)

            word = comtypes.client.CreateObject('Word.Application')
            doc = word.Documents.Open(in_file)
            doc.SaveAs(out_file, FileFormat=wdFormatPDF)
            doc.Close()
            word.Quit()

            os.remove("temp.docx")
        else:
            # Сохранение в word
            doc.save(path[0:len(path)-5] + ".docx") 
    except: return -1
    return 0

if __name__ == "__main__":
    pass